# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT="Noto Sans CJK KR"; NAVY=RGBColor(0x1A,0x4F,0x8B); GREY=RGBColor(0x55,0x55,0x55)
FIG="/sessions/epic-eager-bohr/mnt/스시론 기말 과제/"
doc=Document()
def ea(run,size=10.5,bold=False,color=None,italic=False):
    run.font.name=FONT; run.font.size=Pt(size); run.font.bold=bold; run.font.italic=italic
    if color: run.font.color.rgb=color
    rpr=run._element.get_or_add_rPr(); rf=rpr.get_or_add_rFonts()
    for a in ('w:eastAsia','w:ascii','w:hAnsi'): rf.set(qn(a),FONT)
st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(10.5)
st.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),FONT)
for s in ['Heading 1','Heading 2','Heading 3']:
    hs=doc.styles[s]; hs.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'),FONT)
sec=doc.sections[0]
sec.top_margin=Cm(2); sec.bottom_margin=Cm(2); sec.left_margin=Cm(2.3); sec.right_margin=Cm(2.3)

def H1(t):
    p=doc.add_heading(level=1); r=p.add_run(t); ea(r,15,True,NAVY)
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); return p
def H2(t):
    p=doc.add_heading(level=2); r=p.add_run(t); ea(r,12,True,RGBColor(0x2E,0x5E,0x8C))
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4); return p
def body(t,indent=True):
    p=doc.add_paragraph(); pf=p.paragraph_format
    pf.line_spacing=1.32; pf.space_after=Pt(5)
    if indent: pf.first_line_indent=Cm(0.6)
    r=p.add_run(t); ea(r,10.5); return p
def cap(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); ea(r,9,False,GREY,True); p.paragraph_format.space_after=Pt(8); return p
def img(path,w=15.5):
    doc.add_picture(path, width=Cm(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before=Pt(4)
def lead(topic,rest):
    p=doc.add_paragraph(); pf=p.paragraph_format
    pf.line_spacing=1.32; pf.space_after=Pt(5); pf.first_line_indent=Cm(0.6)
    r=p.add_run(topic); ea(r,10.5,True); r2=p.add_run(" "+rest); ea(r2,10.5); return p
def table(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h); ea(r,9.5,True,RGBColor(0xFF,0xFF,0xFF))
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),'1A4F8B'); c._tc.get_or_add_tcPr().append(sh)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            cs[i].text=''; r=cs[i].paragraphs[0].add_run(str(v)); ea(r,9.5,i==0)
            cs[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
    return t

# ===== 표지 =====
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("수업 프로젝트 보고서"); ea(r,11,False,GREY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30)
r=p.add_run("데이터로 진단하는 업무지구의 성공과 실패"); ea(r,19,True,NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("— 판교테크노밸리와 송도국제업무지구의 비교분석 —"); ea(r,12.5,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(24)
r=p.add_run("가천대학교 스마트시티학과 스마트시티 이론과 실제"); ea(r,11)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("학번 202235870  ·  성다희"); ea(r,11,True)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("분석 시스템: https://jingij2-cmd.github.io/pangyo-songdo-analysis/"); ea(r,9,False,GREY)

# ===== 목차 =====
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(28)
r=p.add_run("목   차"); ea(r,13,True,NAVY)
toc=[("1. 프로젝트 배경 및 목적",["1.1 업무지구 진단의 필요성","1.2 데이터 기반 비교분석 접근","1.3 프로젝트의 목적"]),
 ("2. 비교 대상 선정 및 분석 설계",["2.1 기준지역: 판교 제1테크노밸리","2.2 비교지역: 송도국제업무지구","2.3 분석 범위·데이터·방법"]),
 ("3. 분석 결과",["3.1 토지이용 분석","3.2 교통망 분석 — 등시간권 접근성","3.3 인구사회 및 업무시장 분석"]),
 ("4. 업무지구 성공요인 종합",[]),
 ("5. 결론 및 한계",["5.1 요약 및 시사점","5.2 한계 및 향후 과제"])]
for h,subs in toc:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1); r=p.add_run(h); ea(r,10.5,True)
    for s in subs:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.7); p.paragraph_format.space_after=Pt(1)
        r=p.add_run(s); ea(r,10)
doc.add_page_break()

# ===== 1 =====
H1("1. 프로젝트 배경 및 목적")
H2("1.1 업무지구 진단의 필요성")
lead("동일하게 조성된 업무지구라도 성패는 크게 갈린다.","업무지구는 도시의 일자리와 경제활동이 집중되는 핵심 공간이지만, 물리적 조성을 마쳤다고 모두 성공하는 것은 아니다. 어떤 곳은 기업과 인력이 모이는 거점으로 성장하고, 어떤 곳은 장기간 활성화되지 못한 채 공실과 미개발 부지를 남긴다.")
lead("이 성패의 차이는 입지·교통·토지이용이 복합적으로 작용한 결과이므로, 정성적 인상이 아니라 공공데이터에 근거한 정량적 비교로 진단해야 한다.","특히 업무지구의 경쟁력은 '얼마나 많은 노동시장에 빠르게 접근할 수 있는가'와 '토지가 업무 기능에 얼마나 집중되어 있는가'에 크게 좌우된다. 본 프로젝트는 이 두 축을 데이터로 측정하여 두 업무지구의 성패를 진단한다.")
H2("1.2 데이터 기반 비교분석 접근")
lead("본 연구는 공신력 있는 공공데이터를 결합해 두 업무지구를 동일 기준으로 비교한다.","수도권 지하철 네트워크, 인구·종사자 격자/집계구 통계, 토지이용 용도지역, 오피스 공실률 등을 활용하였다. 또한 분석 결과를 누구나 접속 가능한 웹 기반 등시간권 비교 시스템(GitHub Pages)으로 구축하여, 보고서의 모든 주장이 시스템상의 데이터로 검증되도록 하였다.")
H2("1.3 프로젝트의 목적")
lead("본 프로젝트의 목적은 두 업무지구의 비교를 통해 성공요인을 데이터로 도출하는 것이다.","구체적으로 (1) 성공 업무지구(판교테크노밸리)와 활성화에 실패한 업무지구(송도국제업무지구)를 선정하고, (2) 토지이용·교통망·인구사회 지표를 정량 비교한 뒤, (3) 업무지구의 성공요인을 데이터 근거와 함께 제시한다.")

# ===== 2 =====
H1("2. 비교 대상 선정 및 분석 설계")
H2("2.1 기준지역: 판교 제1테크노밸리")
lead("기준지역은 수도권 대표 성공 업무지구인 판교테크노밸리, 그중 제1판교테크노밸리(약 0.81㎢)로 한정한다.","제2판교테크노밸리는 조성과 입주가 진행 중이어서, '성숙기 업무지구 간 비교'라는 분석의 정합성을 위해 입주·고용이 안정화된 제1판교로 범위를 좁혔다. 핵심역은 신분당선·경강선 환승역인 판교역으로 정의한다.")
H2("2.2 비교지역: 송도국제업무지구 선정 근거")
lead("비교지역으로는 인천 송도국제업무지구(IBD, 약 5.77㎢)를 선정하였으며, 그 근거는 세 가지다.","과제가 요구하는 비교지역의 선정 조건(수도권 내 입지, 조성 완료, 실패 근거의 데이터 입증)을 모두 충족하기 때문이다.")
lead("첫째, 수도권 내에 위치해 동일한 지하철 네트워크로 등시간권을 비교할 수 있다.","송도는 인천1호선 국제업무지구역을 핵심역으로 하여, 판교와 같은 수도권 지하철 그래프 위에서 접근성을 직접 비교할 수 있다.")
lead("둘째, 물리적 조성이 완료되어 비교 가능한 데이터가 존재한다.","송도 IBD는 2002년 사업에 착수해 2011년 지구단위계획이 고시된 대규모 복합 신도시로, 건축물·종사자·토지이용 데이터가 축적되어 있다. 사업이 무산된 지역(예: 용산국제업무지구)과 달리 비교 분석이 가능하다.")
lead("셋째, '활성화 실패'의 근거가 공신력 있는 데이터로 입증된다.","지정 23년이 지났음에도 인천시의회 자료 기준 업무·상업시설 개발률이 2025년 47%(아파트 93%)에 그쳤고, 2022년 1만 명 규모 오피스 유치가 무산되었으며, 시민단체·언론은 이를 '사실상 개발 실패'로 평가한다. 따라서 '성공한 판교'와 대비되는 '저조한 업무지구'의 조건을 명확히 만족한다.")
table(["구분","판교 제1테크노밸리","송도국제업무지구"],
 [["위치","경기 성남시 분당구","인천 연수구 송도동"],
  ["구역 면적","약 0.81㎢","약 5.77㎢"],
  ["핵심역","판교역(신분당·경강선)","국제업무지구역(인천1호선)"],
  ["조성·지정","2011~2015 준공","2011 지구단위계획 고시"]])
cap("[표 1] 두 업무지구의 기본 제원")
H2("2.3 분석 범위·데이터·방법")
lead("세 영역(토지이용·교통·인구사회)을 동일한 시간·공간 기준으로 비교하였다.","시간범위는 인구격자 2024년, 집계구 종사자 2023년, 용도지역 고시연도, 오피스 공실률 2024년 3분기~2026년 1분기이며, 공간범위는 두 구역계 및 핵심역 기준 30·60분 등시간권, 공간단위는 100m 인구격자·집계구·용도지역 폴리곤·지하철 역이다.")
body("등시간권은 지하철 네트워크 그래프에서 핵심역 기준 다익스트라 최단시간(환승 대기 포함)을 구한 뒤, 각 역에서 보행(1.2m/s, 도로우회 1.3배)으로 도달 가능한 범위를 SGIS 100m 인구격자(수도권 25,285,464명) 및 집계구 종사자(수도권 약 1,256만 명)와 공간 결합하여 산출하였다. 토지이용은 VWorld 용도지역(LT_C_UQ111)을 구역계로 클립해 분석하였다. 좌표계 변환(EPSG:5179·5174 ↔ WGS84)은 직접 구현했으며, 송도 구역 산출면적이 경계 속성치의 99.5%로 일치함을 검증하였다.")

# ===== 3 =====
H1("3. 분석 결과")
H2("3.1 토지이용 분석")
lead("판교는 토지가 업무 기능에 고도로 집중된 단일용도 구조다.","구역의 87.4%가 준주거지역으로, 고밀 업무·연구개발 시설이 입지할 수 있는 용도에 토지가 일관되게 배정되어 있다. 나머지는 대부분 녹지(12.6%)로, 업무와 무관한 주거·상업 용도는 사실상 없다.")
lead("반면 송도는 주거·상업·녹지가 분산된 혼합 구조로, 업무 핵심기능에 토지를 집중하지 못했다.","주거가 36.9%, 상업 30.4%, 녹지 32.7%로 세 용도가 비슷한 비중을 차지하며, 특히 자연녹지가 약 3분의 1에 달해 미개발 성격의 토지가 많다. 토지이용 혼합도(엔트로피 LUM)는 판교 0.27, 송도 0.79로, 송도가 업무지구로서 미성숙한 토지이용 구조임을 정량적으로 보여준다.")
table(["용도지역(대분류)","판교 제1테크노밸리","송도 IBD"],
 [["주거","87.4%","36.9%"],["상업","0.0%","30.4%"],["녹지","12.6%","32.7%"],["혼합도 LUM","0.27","0.79"]])
cap("[표 2] 용도지역 구성비 및 토지이용 혼합도")
img(FIG+"fig_landuse.png", 11)
cap("[그림 1] 용도지역 구성비 비교 — 판교는 주거(준주거) 단일용도에 집중, 송도는 분산")
lead("개발 실현 정도에서도 두 지역의 격차가 뚜렷하다.","판교 제1테크노밸리는 계획 용지가 사실상 모두 건축·입주를 마친 반면, 송도 IBD는 업무·상업시설 개발률이 2025년 기준 47%에 그쳐 절반 이상이 미개발·저밀 상태로 남아 있다. 구역의 약 33%가 자연녹지 등 미개발 성격의 토지라는 점도 같은 맥락으로, 송도는 부지는 조성되었으나 토지가 업무 기능으로 충분히 활용되지 못했다. (필지별 정밀 용적률은 건축물대장 자료 확보의 한계로 개발률·미개발지 비율로 갈음하였다.)")
H2("3.2 교통망 분석 — 등시간권 접근성")
lead("판교는 핵심역 기준 30분 등시간권에서 송도를 3배 이상 앞선다.","판교역에서 지하철로 30분 내에 80개 역, 약 105만 명의 인구와 113만 명의 종사자에 도달하는 반면, 송도 국제업무지구역에서는 24개 역, 약 31만 명 인구·14만 명 종사자에 그친다. 60분 기준으로도 판교(788만 인구·541만 종사자)가 송도(250만·100만)를 크게 앞선다.")
table(["30분 등시간권","판교","송도","격차"],
 [["도달역 수","80개","24개","3.3배"],["도달 인구","약 105만","약 31만","3.4배"],["도달 종사자","약 113만","약 14만","8.2배"]])
cap("[표 3] 30분 등시간권 내 도달 가능 인구·종사자")
lead("특히 종사자 격차(8.2배)가 인구 격차(3.4배)보다 훨씬 크다는 점이 핵심이다.","이는 수도권의 일자리가 서울 도심·강남 업무축에 고도로 집중되어 있고, 판교만이 신분당선·경강선 환승을 통해 30분 내 이 노동시장에 직결되기 때문이다. 송도는 인천1호선 단일노선에 의존해 등시간권이 인천 내부에 갇히며 서울 도심에 도달하지 못한다.")
img(FIG+"fig_isochrone_map.png", 13.5)
cap("[그림 2] 핵심역 기준 30·60분 등시간권 도달 인구격자 — 판교는 서울 전역, 송도는 인천 내부로 한정")
img(FIG+"fig_accessibility_curve.png", 13.5)
cap("[그림 3] 누적 접근성 곡선 — 소요시간별 도달 가능 인구·종사자. 모든 구간에서 판교가 송도를 크게 상회")
H2("3.3 인구사회 및 업무시장 분석")
lead("등시간권 도달 종사자는 곧 각 업무지구가 동원할 수 있는 노동시장의 규모를 의미한다.","판교는 30분 내 약 113만, 60분 내 약 541만 명의 종사자 풀에 접근 가능한 반면 송도는 각각 14만·100만 명에 불과하다. 우수 인력 접근성이 업무지구 성패의 핵심 자원이라는 점에서, 이 격차는 판교의 기업 집적과 송도의 기업 유치 부진을 구조적으로 설명한다.")
lead("이러한 접근성 우위는 실제 시장 성과로 확인된다.","판교(제1·2 합산)는 2024년 입주기업 1,803개사(전년 대비 11%↑), 임직원 약 7만 9천 명, 2022년 매출 167.7조 원을 기록하며 IT 중심 첨단산업 클러스터로 성장했다. 반면 송도는 대형 오피스 유치가 무산되고 업무 개발률이 절반에 못 미쳤으며, 오피스 공실률에서도 인천 권역이 16.9%(2026년 1분기, 최고 20.8%)로 전국 평균 8.8%의 약 두 배에 달해 업무시장 수요 부진을 드러낸다.")
img(FIG+"fig_vacancy.png", 11.5)
cap("[그림 4] 오피스 공실률 추이 — 인천 권역은 전국 평균의 약 2배로 업무시장 수요 부진을 보여줌")

# ===== 4 =====
H1("4. 업무지구 성공요인 종합")
body("두 업무지구의 성패는 물리적 조성 여부가 아니라 세 가지 요인으로 일관되게 설명된다. 분석 결과를 근거 단위로 정리하면 다음과 같다.")
lead("첫째, 판교는 노동시장 접근성에서 구조적 우위를 가진다.","대중교통으로 30분 내 도달 가능한 종사자가 판교는 약 113만 명인 반면 송도는 14만 명으로 8배 이상 차이가 난다. 판교는 복수 노선 환승과 서울 도심 직결로 거대한 인력 풀에 닿지만, 송도는 단일노선 의존으로 인력 확보에서 근본적 한계를 안고 있다.")
lead("둘째, 판교는 토지이용이 업무 기능에 집중되어 있다.","판교는 구역의 87%가 준주거 단일용도로 업무·연구에 토지를 집중한 반면, 송도는 주거·상업·녹지가 고르게 혼재(혼합도 0.79)되고 자연녹지가 3분의 1을 차지해 업무 핵심기능이 분산·미성숙하다. 즉 판교는 '업무에 특화된 계획', 송도는 '분산된 계획'을 가진 셈이다.")
lead("셋째, 이러한 입지·계획의 우위가 실제 시장 성과로 귀결되었다.","판교는 입주기업·종사자·매출이 지속 성장하며 활성화에 성공한 반면, 송도는 업무 개발률 47%·높은 공실률·오피스 유치 무산으로 활성화에 실패하였다. 접근성과 토지이용이라는 원인이 시장 성과라는 결과로 이어진 것이다.")
body("요컨대 업무지구의 성공은 부지를 조성했는가가 아니라, '대규모 노동시장에 대한 접근성'과 '업무 기능에 집중된 토지이용'을 갖추었는가에 의해 결정된다. 판교는 두 조건을 모두 충족했고, 송도는 두 조건 모두에서 취약하여 동일한 물리적 조성에도 상반된 결과를 낳았다.")

# ===== 5 =====
H1("5. 결론 및 한계")
H2("5.1 요약 및 시사점")
lead("본 프로젝트는 업무지구 성공의 핵심이 '노동시장 접근성'과 '업무 중심 토지이용'임을 데이터로 입증하였다.","판교와 송도를 토지이용·교통·인구사회 지표로 정량 비교한 결과, 두 조건을 충족한 판교는 성공하고 그렇지 못한 송도는 실패하는 일관된 패턴이 확인되었다.")
lead("이는 신규 업무단지 조성에서 부지 공급보다 접근성과 토지이용 계획이 선행되어야 함을 시사한다.","송도의 사례는 광역철도 연결 강화와 업무용지 중심의 토지이용 재편 없이는 활성화가 어렵다는 정책적 함의를 제공한다.")
H2("5.2 한계 및 향후 과제")
body("본 분석은 다음 한계를 가진다. 첫째, 판교 종사자·매출 공식 통계는 제1·2판교 합산 위주여서 제1판교 단독 수치는 일부 추정을 포함한다. 둘째, 오피스 공실률은 송도·판교 단독 상권이 없어 각각 인천 권역·분당역세권을 프록시로 사용하였다. 셋째, 등시간권 보행망은 직선거리에 도로우회계수(1.3)를 적용한 근사이며 실제 보행 네트워크와 차이가 있을 수 있다. 넷째, 두 구역의 면적 규모 차이(0.81㎢ vs 5.77㎢)가 커 절대값보다 밀도·비율·접근성 지표 중심으로 해석하였다. 향후 실제 보행 네트워크와 시계열 데이터를 결합하면 조성 전후의 변화까지 분석할 수 있을 것이다.")
body("[데이터 출처] SGIS 통계지리정보서비스(인구격자·집계구 종사자), VWorld 토지이용(LT_C_UQ111), 한국부동산원 R-ONE 상업용부동산 임대동향, 경기도 2024 판교테크노밸리 실태조사, 인천시의회·경인미래신문·경기일보 보도, 수도권 지하철 네트워크 그래프(수업 제공).", indent=False)

import os
out="/sessions/epic-eager-bohr/mnt/스시론 기말 과제/판교_송도_분석보고서.docx"
doc.save(out); print("saved",out,os.path.getsize(out))
